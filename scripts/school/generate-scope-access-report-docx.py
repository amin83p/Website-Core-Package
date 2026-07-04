#!/usr/bin/env python3
"""Generate School Package Scope Access Report as Word document."""

import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'python-docx', '-q'])
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parents[2]
OUTPUT = ROOT / 'docs' / 'school-scope-access-report.docx'


def set_doc_defaults(doc):
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)


def add_title(doc, text):
    p = doc.add_heading(text, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_h1(doc, text):
    doc.add_heading(text, level=1)


def add_h2(doc, text):
    doc.add_heading(text, level=2)


def add_h3(doc, text):
    doc.add_heading(text, level=3)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style='List Bullet')


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
    doc.add_paragraph('')


def build_document():
    doc = Document()
    set_doc_defaults(doc)

    add_title(doc, 'School Package Scope Access Report')
    add_para(doc, 'Website Core Package — School Module')
    add_para(doc, 'Generated: July 2026')
    add_para(doc, 'This report documents how access-profile scopes affect data visibility and operations across school package sections.')
    doc.add_page_break()

    add_h1(doc, '1. How School Access Works')
    add_para(doc, 'Every request is checked in two layers:')
    add_bullets(doc, [
        'Section + operation (access profile): Route middleware (requireAccess) decides whether the user may open that area. Without the operation on that section → 403.',
        'Scope (SCP_ADMIN, SCP_ORG, SCP_DEPT, SCP_DIV, SCP_OWNER, SCP_USER): After the route allows access, scope limits which records in the active org are visible or mutable.',
    ])

    add_h2(doc, '1.1 Scope definitions (school domain)')
    add_table(doc,
        ['Scope', 'Data visibility'],
        [
            ['ADMIN', 'All records in the active org, no limitation'],
            ['ORGANIZATION', 'Same as ADMIN in school — all org data for granted sections'],
            ['DEPARTMENT / DIVISION', 'Records tied to the user\'s linked person (instructor, assignee, deliveredBy, etc.)'],
            ['OWNER', 'Only records the user created (ownerUserId, audit.createUser, etc.)'],
            ['USER', 'No school data (empty lists / not found)'],
        ])

    add_h2(doc, '1.2 Legend for tables')
    add_table(doc,
        ['Symbol', 'Meaning'],
        [
            ['All', 'All records in active org'],
            ['Assigned', 'Person-assignment filter (instructor / assignee / deliveredBy)'],
            ['Own', 'Creator-only records'],
            ['Own person', 'Only the user\'s linked person row'],
            ['None', 'No data'],
            ['Route only', 'Scope not passed from controller; defaults to org-wide reads if route allows'],
            ['Custom', 'Section-specific logic, not req.accessScope'],
            ['Session†', 'Manage Session gate: DEPT/DIV = must be deliveredBy; OWNER = must be session creator'],
            ['Partial', 'Scope infrastructure exists but not fully wired to all endpoints'],
        ])

    add_h2(doc, '1.3 Implementation status summary')
    add_table(doc,
        ['Section group', 'Scope wired?', 'Primary enforcement'],
        [
            ['Classes, sessions, enrollment', 'Yes', 'req.accessScope + schoolRecordAccessService'],
            ['Activities', 'Yes', 'Assignee filter in repository'],
            ['Schedule events', 'Partial', 'accessContext on event builder + UI person lock'],
            ['Students', 'Yes', 'personId + owner OR via buildRouteAccessContext'],
            ['Teachers, staff', 'Yes', 'personId + owner OR via buildRouteAccessContext'],
            ['School accounts', 'Yes', 'partyAccounts linked IDs + owner OR'],
            ['Timesheets (list)', 'Partial', 'Scope + teacher/admin context'],
            ['Master Academia Hub', 'Partial', 'scopeId on some hub modules'],
            ['Programs, catalog data', 'Partial', 'Catalog kinds org-visible for labels'],
            ['Reports, exams', 'No', 'Custom admin + person/assignment rules'],
            ['Leave, tasks', 'No', 'Custom own-record + section admin'],
            ['Finance, registrations, ledger', 'No', 'Route only'],
        ])

    doc.add_page_break()

    add_h1(doc, '2. All School Package Sections')
    add_table(doc,
        ['Section ID', 'UI area'],
        [
            ['SCHOOL', 'School dashboard'],
            ['SCHOOL_MASTER_ACADEMIA_HUB', 'Master Academia Hub'],
            ['SCHOOL_DEPARTMENTS', 'Departments'],
            ['SCHOOL_TERMS', 'Terms'],
            ['SCHOOL_PROGRAMS', 'Programs'],
            ['SCHOOL_SUBJECTS', 'Subjects'],
            ['SCHOOL_SESSION_STATUSES', 'Session statuses'],
            ['SCHOOL_HOLIDAYS', 'Holidays'],
            ['SCHOOL_CLASSES', 'Classes'],
            ['SCHOOL_CLASS_ENROLLMENT_PERIODS', 'Rolling enrollment periods'],
            ['SCHOOL_CLASS_CYCLES', 'Class cycles / rollover'],
            ['SCHOOL_SESSIONS', 'Manage Session, attendance, gradebook, cases'],
            ['SCHOOL_GRADEBOOK', 'Gradebook / final grades'],
            ['SCHOOL_ATTENDANCES', 'Attendance matrix'],
            ['SCHOOL_STUDENTS', 'Students'],
            ['SCHOOL_TEACHERS', 'Teachers'],
            ['SCHOOL_STAFF', 'Staff'],
            ['SCHOOL_PAY_RATES', 'Pay rates'],
            ['SCHOOL_SCHEDULES', 'Schedule viewer'],
            ['SCHOOL_CALENDAR', 'School calendar'],
            ['SCHOOL_ACTIVITIES', 'Activities'],
            ['SCHOOL_TASKS', 'Tasks'],
            ['SCHOOL_LEAVE_REQUESTS', 'Leave requests'],
            ['SCHOOL_REPORTS', 'Reports hub'],
            ['SCHOOL_REPORTS_TEMPLATE', 'Report templates'],
            ['SCHOOL_REPORTS_ASSIGNMENT', 'Report assignments'],
            ['SCHOOL_REPORTS_INSTANCES', 'Report instances'],
            ['SCHOOL_EXAMS', 'Exams hub'],
            ['SCHOOL_EXAMS_TEMPLATE', 'Exam templates'],
            ['SCHOOL_EXAMS_ALLOCATION', 'Exam allocations'],
            ['SCHOOL_EXAMS_TAKING', 'Exam taking'],
            ['SCHOOL_EXAMS_REVIEW', 'Exam review'],
            ['SCHOOL_TIMESHEET_PERIODS', 'Timesheet periods'],
            ['SCHOOL_TIMESHEETS', 'Timesheets'],
            ['SCHOOL_TIMESHEET_MANAGEMENT', 'Timesheet management'],
            ['SCHOOL_PROGRAM_REGISTRATIONS', 'Program registrations'],
            ['SCHOOL_TERM_REGISTRATIONS', 'Term registrations'],
            ['SCHOOL_PRIOR_SUBJECT_CREDITS', 'Prior subject credits'],
            ['SCHOOL_ACADEMIC_LEDGER', 'Academic ledger'],
            ['SCHOOL_WITHDRAWAL', 'Withdrawals'],
            ['SCHOOL_TRANSACTION_TEMPLATES', 'Fee / transaction definitions'],
            ['SCHOOL_ACCOUNTS', 'School accounts'],
            ['SCHOOL_TRANSACTIONS', 'Transactions manager'],
            ['SCHOOL_SAMPLE_DATA', 'Sample data tools'],
        ])

    doc.add_page_break()

    add_h1(doc, '3. Part 1 — Master Data & Classes / Sessions')
    add_para(doc, 'Operations: R = READ/READ_ALL, C = CREATE, U = UPDATE, D = DELETE. All require the operation on the section in the access profile first.')

    sections_part1 = [
        ('Departments, Terms, Session statuses, Holidays, Subjects (catalog)',
         'Catalog entities: org-wide read under assignment scope for dropdowns.',
         [['R', 'All', 'All (catalog)', 'Own', 'None', 'Route only'],
          ['C', 'All', 'If route grants', 'Yes', 'None', 'Route only'],
          ['U', 'All', 'All if reachable', 'Own', 'None', 'Route only'],
          ['D', 'All', 'All if reachable', 'Own', 'None', 'Route only']]),
        ('Programs',
         'No specialized assignment filter; unwired controllers default to org-wide.',
         [['R', 'All', 'Route only (all org)', 'Own', 'None', 'Route only'],
          ['C/U/D', 'All', 'Same as R', 'Own', 'None', 'Route only']]),
        ('Classes (SCHOOL_CLASSES)',
         'Scope fully wired via buildRouteAccessContext.',
         [['R', 'All', 'Instructor classes', 'Own created', 'None', 'Yes'],
          ['C', 'Yes', 'Yes', 'Yes', 'None', 'Yes'],
          ['U', 'Any class', 'Instructor classes', 'Own only', 'None', 'Yes'],
          ['D', 'Any class', 'Instructor classes', 'Own only', 'None', 'Yes']]),
        ('Class enrollment / cycles',
         'Parent class must be in scope.',
         [['R/U/C/D', 'All', 'Same as Classes', 'Same as Classes', 'None', 'Yes']]),
        ('Sessions & Manage Session (SCHOOL_SESSIONS)',
         'Strict session† gates on Manage Session and mutations.',
         [['R (list)', 'All', 'Instructor or deliveredBy', 'Own classes/sessions', 'None', 'Yes'],
          ['R (Manage Session)', 'Any session', 'Session† deliveredBy', 'Session† creator', 'None', 'Yes'],
          ['U (gradebook, cases, etc.)', 'Any', 'Session†', 'Session†', 'None', 'Yes'],
          ['C (makeup)', 'Yes', 'Session†', 'Session†', 'None', 'Yes']]),
        ('Gradebook / final grades (SCHOOL_GRADEBOOK)',
         'Session gradebook uses session† rules; final grades use class scope.',
         [['R/U (session)', 'All', 'Session†', 'Session†', 'None', 'Yes'],
          ['R/U (final grades)', 'All', 'Instructor class', 'Own class', 'None', 'Partial']]),
    ]

    for title, note, rows in sections_part1:
        add_h2(doc, title)
        add_para(doc, note)
        add_table(doc, ['Op', 'ADMIN/ORG', 'DEPT/DIV', 'OWNER', 'USER', 'Wired?'], rows)

    doc.add_page_break()

    add_h1(doc, '4. Part 2 — People, Scheduling, Activities, Leave & Tasks')

    sections_part2 = [
        ('Students (SCHOOL_STUDENTS)',
         'Division/Dept: own person row OR records user created (personId + owner OR).',
         [['R', 'All', 'Own person OR created', 'Own created', 'None', 'Yes'],
          ['C/U/D', 'All', 'If reachable', 'Own', 'None', 'Partial']]),
        ('Teachers (SCHOOL_TEACHERS)',
         'Division/Dept: own teacher row OR records user created.',
         [['R', 'All', 'Own teacher row OR created', 'Own created', 'None', 'Yes'],
          ['C/U/D', 'All', 'Own row if reachable', 'Own', 'None', 'Partial']]),
        ('Staff (SCHOOL_STAFF)', 'Same as Teachers.', [['R', 'All', 'Own staff row OR created', 'Own created', 'None', 'Yes'], ['C/U/D', 'All', 'Own if reachable', 'Own', 'None', 'Partial']]),
        ('Pay rates (SCHOOL_PAY_RATES)', 'Controller does not pass scope.', [['R/C/U/D', 'All', 'Route only', 'Own', 'None', 'Route only']]),
        ('Schedules (SCHOOL_SCHEDULES)',
         'Non-admins locked to own personId; event builder passes accessContext.',
         [['R (my schedule)', 'All events', 'Scoped classes + own person', 'Own person', 'None', 'Partial'],
          ['R (other person)', 'Schedule admin', 'Denied', 'Denied', 'None', 'Custom UI']]),
        ('School calendar (SCHOOL_CALENDAR)', 'Aggregator; calendar admin can pick persons.', [['R', 'All layers', 'Locked to self', 'Own person', 'None', 'Custom UI'], ['C/U/D', 'N/A', 'N/A', 'N/A', 'N/A', '—']]),
        ('Holidays (SCHOOL_HOLIDAYS)', 'Catalog in repository.', [['R/C/U/D', 'All', 'All (catalog)', 'Own', 'None', 'Route only']]),
        ('Activities (SCHOOL_ACTIVITIES)', 'Assignee filter wired.', [['R', 'All', 'Assignee', 'Own', 'None', 'Yes'], ['C/U/D', 'All', 'Assigned', 'Own', 'None', 'Yes']]),
        ('Leave requests (SCHOOL_LEAVE_REQUESTS)',
         'Custom: own requests OR leave section admin. NOT access-profile scope.',
         [['R', 'All (admin)', 'Own requests‡', 'Own requests‡', 'None', 'Custom'],
          ['C/U', 'Admin or own', 'Own', 'Own', 'None', 'Custom'],
          ['Review', 'Section admin', 'Denied', 'Denied', 'Denied', 'Custom']]),
        ('Tasks (SCHOOL_TASKS)',
         'Custom: assigned person/role OR tasks section admin.',
         [['R', 'All (admin)', 'Assigned/role', 'Assigned', 'None', 'Custom'],
          ['U', 'Admin or assignee', 'Assignee', 'Assignee', 'None', 'Custom'],
          ['D', 'Section admin', 'Denied', 'Denied', 'None', 'Custom']]),
    ]

    for title, note, rows in sections_part2:
        add_h2(doc, title)
        add_para(doc, note)
        add_table(doc, ['Op', 'ADMIN/ORG', 'DEPT/DIV', 'OWNER', 'USER', 'Wired?'], rows)

    doc.add_page_break()

    add_h1(doc, '5. Part 3 — Reports, Exams, Timesheets, Finance & Admin')

    sections_part3 = [
        ('Report templates & assignments', 'Route only; no access-profile scope.', [['R/C/U/D', 'All', 'Route only', 'Own', 'None', 'Route only']]),
        ('Report instances',
         'Session rows use canViewerSeeSessionReportRow + session† on Manage Session.',
         [['R (general)', 'All', 'Route only', 'Own', 'None', 'Route only'],
          ['R (session reports)', 'All', 'Custom person/assignment', 'Custom', 'None', 'Custom + session†']]),
        ('Exam templates', 'Route only.', [['R/C/U/D', 'All', 'Route only', 'Own', 'None', 'Route only']]),
        ('Exam allocations & review',
         'Custom: creator OR class instructor OR examAssignments.personId.',
         [['R', 'All', 'Custom assigned', 'Creator/assigned', 'None', 'Custom'],
          ['C/U/D', 'Exam admin/workflow', 'Instructor/assigned', 'Own/assigned', 'None', 'Custom']]),
        ('Exam taking (SCHOOL_EXAMS_TAKING)', 'Linked person on assignment; session context on class pages.', [['START/SAVE', 'Staff/admin', 'Assigned person', 'Own attempts', 'None', 'Custom']]),
        ('Timesheet periods', 'Hub passes scopeId; controller route only.', [['R/C/U/D', 'All', 'Route only', 'Own', 'None', 'Partial']]),
        ('Timesheets (SCHOOL_TIMESHEETS)',
         'Partial scope + resolveTargetTeacherContext (admin can pick teacher).',
         [['R', 'All / pick teacher', 'Own personId rows', 'Own', 'None', 'Partial'],
          ['U', 'Own or admin', 'Own only', 'Own', 'None', 'Custom']]),
        ('Timesheet management', 'Section admin only for approve/reopen.', [['R/U', 'Timesheet admin', 'Denied', 'Denied', 'None', 'Section admin']]),
        ('Attendance matrix (SCHOOL_ATTENDANCES)', 'Class list route only; writes via session†.', [['R', 'All', 'Route only', 'Own classes', 'None', 'Route only'], ['U', 'Via session', 'Session†', 'Session†', 'None', 'Session wired']]),
        ('Programs & registrations', 'Route only on all registration flows.', [['R/C/U', 'All', 'Route only', 'Own', 'None', 'Route only']]),
        ('Academic ledger & withdrawal', 'Route only; ledger is append-only.', [['R', 'All', 'Route only', 'Own', 'None', 'Route only'], ['C', 'Yes', 'Yes', 'Yes', 'None', 'Route only']]),
        ('School accounts (SCHOOL_ACCOUNTS)',
         'Division/Dept: linked teacher/staff/student account OR accounts user created.',
         [['R', 'All', 'Linked party OR created', 'Own created', 'None', 'Yes'],
          ['C/U/D', 'All', 'If reachable', 'Own', 'None', 'Partial']]),
        ('Finance (templates, accounts, transactions)', 'School accounts scoped; transactions still route-only.', [['R/C/U', 'All', 'Accounts: linked+owner', 'Own', 'None', 'Partial']]),
        ('Master Academia Hub', 'Passes scopeId for classes, departments, periods, holidays.', [['Hub data', 'All org', 'Per module (classes scoped)', 'Own where applicable', 'None', 'Partial']]),
        ('Sample data (SCHOOL_SAMPLE_DATA)', 'Admin tooling.', [['All', 'Admin typical', 'Limited', 'Denied', 'None', 'Admin']]),
    ]

    for title, note, rows in sections_part3:
        add_h2(doc, title)
        add_para(doc, note)
        add_table(doc, ['Op', 'ADMIN/ORG', 'DEPT/DIV', 'OWNER', 'USER', 'Wired?'], rows)

    doc.add_page_break()

    add_h1(doc, '6. Key Rules Reference')
    add_h3(doc, 'Manage Session strict gates (implemented)')
    add_bullets(doc, [
        'ADMIN / ORGANIZATION: Any session in the active org (if route grants access).',
        'DEPARTMENT / DIVISION: User\'s personId must match session.delivery.deliveredBy. Class instructor alone is NOT sufficient.',
        'OWNER: User must be the creator of that specific session (audit.createUser, makeup.createdBy), not merely the class creator.',
        'USER: Always denied.',
    ])

    add_h3(doc, 'Files implementing scope (reference)')
    add_bullets(doc, [
        'packages/school/MVC/services/school/schoolDataScopeBuilder.js — scope mode resolution',
        'packages/school/MVC/services/school/schoolRecordAccessService.js — record/session asserts',
        'packages/school/MVC/repositories/school/index.js — list filters by scopeMode',
        'packages/school/MVC/controllers/school/classController.js — Manage Session wiring',
        'packages/school/config/accessConstants.js — section identifiers',
    ])

    add_h3(doc, 'Recommended follow-up')
    add_bullets(doc, [
        'Migrate remaining controllers to pass buildRouteAccessContext(req) on all user-facing fetchData/getDataById calls.',
        'Align exams, reports, leave, and tasks with access-profile scope or document them as intentional custom models.',
        'Fix student list assignment filter to support "students in my classes" for DEPT/DIV teachers.',
        'Backfill ownerUserId / audit.createUser on legacy records for OWNER scope visibility.',
    ])

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    return OUTPUT


if __name__ == '__main__':
    path = build_document()
    print(f'Generated: {path}')
